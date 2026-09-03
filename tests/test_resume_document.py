import io
import unittest
from pathlib import Path

from docx import Document

from src.core.resume_document import MAX_FILE_BYTES, ResumeDocumentError, extract_resume_document


class ResumeDocumentTest(unittest.TestCase):
    def test_txt_extracts_draft_fields(self):
        text = "教育背景\n硕士 人工智能\n工作经历\n3年开发经验，负责Python平台\n项目经历\n使用RAG和Docker\n专业技能\nPython、RAG、Docker"
        result = extract_resume_document("简历.txt", text.encode())
        self.assertEqual(result.file_type, "txt")
        self.assertIn("硕士", result.education)
        self.assertIn("3年", result.experience)
        self.assertIn("RAG", result.projects)
        self.assertIn("Docker", result.skills_raw)

    def test_docx_extracts_paragraphs_and_tables(self):
        document = Document(); document.add_paragraph("专业技能"); document.add_paragraph("Python、LangGraph")
        table = document.add_table(rows=1, cols=2); table.cell(0, 0).text = "项目经历"; table.cell(0, 1).text = "RAG知识库"
        stream = io.BytesIO(); document.save(stream)
        result = extract_resume_document("resume.docx", stream.getvalue())
        self.assertIn("Python", result.raw_text); self.assertIn("RAG知识库", result.raw_text)

    def test_rejects_old_doc_spoofed_pdf_binary_and_oversize(self):
        cases = [("resume.doc", b"legacy"), ("resume.pdf", b"not-pdf"),
                 ("resume.txt", b"a" * (MAX_FILE_BYTES + 1))]
        for name, data in cases:
            with self.subTest(name=name), self.assertRaises(ResumeDocumentError):
                extract_resume_document(name, data)

    def test_filename_is_reduced_to_basename(self):
        self.assertEqual(extract_resume_document("../../private/resume.txt", b"Python").file_name, "resume.txt")

    def test_extracts_identity_contact_target_major_and_bounded_sections(self):
        text = """张三
电话：13800138000  邮箱：zhangsan@example.com
求职意向：AI Agent开发工程师
教育背景
某大学 人工智能专业 硕士
工作经验：1 年（项目实践与实习经历）
工作经历
2025.01-2025.12 某公司实习，使用 Python/FastAPI 开发接口。
项目经历
基于 LangGraph 设计 Agent 工作流。
专业技能
Python、FastAPI、LangGraph、MCP
"""
        result = extract_resume_document("张三.txt", text.encode())
        self.assertEqual(result.name, "张三")
        self.assertEqual(result.phone, "13800138000")
        self.assertEqual(result.email, "zhangsan@example.com")
        self.assertEqual(result.target_job, "AI Agent开发工程师")
        self.assertEqual(result.target_job_source, "explicit")
        self.assertEqual(result.degree, "硕士")
        self.assertEqual(result.major, "人工智能")
        self.assertEqual(result.experience, "1 年（项目实践与实习经历）")
        self.assertEqual(result.experience_source, "explicit")
        self.assertIn("某公司实习", result.work_experience)
        self.assertNotIn("LangGraph", result.work_experience)
        self.assertIn("LangGraph", result.projects)
        self.assertNotIn("专业技能", result.projects)

    def test_explicit_experience_variants_take_priority(self):
        for value in ("1年", "1 年", "工作经验：1年", "1年工作经验"):
            with self.subTest(value=value):
                result = extract_resume_document("resume.txt", f"{value}\n工作经历\n某公司实习".encode())
                self.assertIn("1", result.experience)
                self.assertEqual(result.experience_source, "explicit")

    def test_infers_experience_only_from_work_date_ranges(self):
        text = "工作经历\n2024.01-2024.12 某公司实习\n项目经历\n2022.01-2025.12 学生项目"
        result = extract_resume_document("resume.txt", text.encode())
        self.assertRegex(result.experience, r"约1(?:\.0)?年")
        self.assertEqual(result.experience_source, "date_range_inferred")

    def test_infers_target_job_conservatively_from_existing_resume_content(self):
        data = (Path(__file__).parent / "fixtures" / "sample_resume.txt").read_bytes()
        result = extract_resume_document("sample_resume.txt", data)
        self.assertEqual(result.target_job, "AI Agent开发工程师")
        self.assertEqual(result.target_job_source, "content_inferred")


if __name__ == "__main__":
    unittest.main()
