"""Safe, in-memory text extraction for user-supplied resumes."""
from __future__ import annotations

import io
import re
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from ..schemas import ResumeDocumentExtractResult

MAX_FILE_BYTES = 8 * 1024 * 1024
MAX_TEXT_CHARS = 200_000
MAX_PDF_PAGES = 100
MAX_DOCX_UNCOMPRESSED = 32 * 1024 * 1024
MAX_DOCX_ENTRIES = 1000
SUPPORTED = {".pdf": "pdf", ".docx": "docx", ".txt": "txt"}


class ResumeDocumentError(ValueError):
    pass


def _clean_text(value: str) -> str:
    value = value.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    value = "".join(char for char in value if char in "\n\t" or ord(char) >= 32)
    value = re.sub(r"[ \t]+", " ", value)
    return re.sub(r"\n{3,}", "\n\n", value).strip()[:MAX_TEXT_CHARS]


def _extract_pdf(data: bytes) -> tuple[str, list[str]]:
    if not data.startswith(b"%PDF-"):
        raise ResumeDocumentError("文件扩展名是 PDF，但内容不是有效 PDF。")
    try:
        reader = PdfReader(io.BytesIO(data), strict=True)
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ResumeDocumentError(f"PDF 页数不能超过 {MAX_PDF_PAGES} 页。")
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except ResumeDocumentError:
        raise
    except Exception as exc:
        raise ResumeDocumentError("PDF 无法解析，可能已损坏或加密。") from exc
    warnings = [] if text.strip() else ["该 PDF 没有可提取的文本层，可能是扫描件；请使用 OCR 后重试或手工粘贴文本。"]
    return text, warnings


def _extract_docx(data: bytes) -> tuple[str, list[str]]:
    if not data.startswith(b"PK"):
        raise ResumeDocumentError("文件扩展名是 DOCX，但内容不是有效 Word 文档。")
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_DOCX_ENTRIES or sum(item.file_size for item in infos) > MAX_DOCX_UNCOMPRESSED:
                raise ResumeDocumentError("DOCX 解压后体积异常，已拒绝解析。")
            if "word/document.xml" not in archive.namelist():
                raise ResumeDocumentError("DOCX 缺少正文结构。")
        document = Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                line = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if line:
                    parts.append(line)
        return "\n".join(parts), []
    except ResumeDocumentError:
        raise
    except Exception as exc:
        raise ResumeDocumentError("DOCX 无法解析，可能已损坏或包含不支持的结构。") from exc


def _extract_txt(data: bytes) -> tuple[str, list[str]]:
    if b"\x00" in data:
        raise ResumeDocumentError("TXT 中包含二进制内容，已拒绝解析。")
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding), []
        except UnicodeDecodeError:
            continue
    raise ResumeDocumentError("TXT 编码无法识别，请保存为 UTF-8 后重试。")


HEADINGS = {
    "education": ("教育经历", "教育背景", "学历信息"),
    "work_experience": ("工作经历", "工作经验", "职业经历", "实习经历"),
    "projects": ("项目经历", "项目经验", "代表项目"),
    "skills_raw": ("专业技能", "技能清单", "技能特长", "核心技能", "个人技能"),
}

OTHER_HEADINGS = {
    "基本信息", "个人信息", "求职意向", "目标岗位", "联系方式", "证书", "证书荣誉",
    "荣誉奖项", "校园经历", "自我评价", "个人评价", "语言能力", "培训经历",
}


def _heading(line: str) -> tuple[str | None, str] | None:
    compact = re.sub(r"\s+", "", line).strip()
    for field, names in HEADINGS.items():
        for name in names:
            if compact == name:
                return field, ""
            match = re.match(rf"^{re.escape(name)}[：:](.*)$", compact)
            if match:
                return field, match.group(1).strip()
    for name in OTHER_HEADINGS:
        if compact == name or re.match(rf"^{re.escape(name)}[：:]", compact):
            return None, ""
    return None


def _sections(lines: list[str]) -> dict[str, list[str]]:
    sections = {field: [] for field in HEADINGS}
    active: str | None = None
    for line in lines:
        heading = _heading(line)
        if heading is not None:
            active, inline = heading
            if active and inline:
                sections[active].append(inline)
            continue
        if active:
            sections[active].append(line)
    return sections


def _explicit_experience(lines: list[str]) -> str:
    patterns = (
        r"工作经验\s*[：:]\s*(\d{1,2}(?:\.\d+)?\s*年(?:以上)?(?:\s*[（(][^）)\n]{1,40}[）)])?)",
        r"(?<!\d)(\d{1,2}(?:\.\d+)?\s*年(?:以上)?(?:\s*[（(][^）)\n]{1,40}[）)])?)\s*(?:工作|开发|相关|实习|项目)?经验",
        r"^\s*(\d{1,2}(?:\.\d+)?\s*年(?:以上)?(?:\s*[（(][^）)\n]{1,40}[）)])?)\s*$",
    )
    for line in lines:
        for pattern in patterns:
            if match := re.search(pattern, line):
                return match.group(1).strip()
    return ""


DATE_RANGE = re.compile(
    r"(?P<sy>20\d{2})[./年-](?P<sm>\d{1,2})月?\s*(?:-|—|–|~|至|到)\s*"
    r"(?:(?P<ey>20\d{2})[./年-](?P<em>\d{1,2})月?|(?P<present>至今|现在|今))"
)


def _inferred_experience(work_text: str) -> str:
    intervals: list[tuple[int, int]] = []
    today = date.today()
    for match in DATE_RANGE.finditer(work_text):
        start = int(match.group("sy")) * 12 + int(match.group("sm")) - 1
        if match.group("present"):
            end = today.year * 12 + today.month - 1
        else:
            end = int(match.group("ey")) * 12 + int(match.group("em")) - 1
        if 0 <= end - start <= 600:
            intervals.append((start, end + 1))
    if not intervals:
        return ""
    merged: list[list[int]] = []
    for start, end in sorted(intervals):
        if merged and start <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], end)
        else:
            merged.append([start, end])
    months = sum(end - start for start, end in merged)
    if months < 12:
        return f"约{months}个月（根据工作/实习时间范围推算）"
    years = months / 12
    value = str(int(years)) if years.is_integer() else f"{years:.1f}"
    return f"约{value}年（根据工作/实习时间范围推算）"


def _first_match(pattern: str, text: str, flags: int = 0) -> str:
    match = re.search(pattern, text, flags)
    return match.group(1).strip() if match else ""


def _infer_target_job(text: str, filename: str) -> tuple[str, str]:
    filename_text = re.sub(r"[_-]+", " ", Path(filename).stem)
    rules = (
        ("AI Agent开发工程师", r"AI\s*Agent|Agent开发|智能体"),
        ("RAG引擎研发工程师", r"RAG|检索增强"),
        ("AI安全技术工程师", r"AI安全|LLM安全|提示注入|Guardrails?"),
    )
    for job, pattern in rules:
        if re.search(pattern, filename_text, re.IGNORECASE):
            return job, "filename_inferred"
    if re.search(r"(?:LangGraph|MCP)", text, re.IGNORECASE) and re.search(r"Agent|智能体", text, re.IGNORECASE):
        return "AI Agent开发工程师", "content_inferred"
    if re.search(r"AI安全|LLM安全|提示注入|Guardrails?", text, re.IGNORECASE):
        return "AI安全技术工程师", "content_inferred"
    if re.search(r"RAG|检索增强", text, re.IGNORECASE) and re.search(r"检索|知识库|向量", text):
        return "RAG引擎研发工程师", "content_inferred"
    return "", "unknown"


def _draft_fields(text: str, filename: str = "") -> dict[str, str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    sections = _sections(lines)
    education = "\n".join(sections["education"])
    if not education:
        education = next((line for line in lines if re.search(r"博士|硕士|本科|大专|专科", line)), "")
    work = "\n".join(sections["work_experience"])
    projects = "\n".join(sections["projects"])
    skills = "\n".join(sections["skills_raw"])
    if not any((work, projects, skills)):
        work = text
    experience = _explicit_experience(lines)
    experience_source = "explicit" if experience else "unknown"
    if not experience:
        experience = _inferred_experience(work)
        experience_source = "date_range_inferred" if experience else "unknown"
    if not experience:
        experience = next(("应届毕业生" for line in lines if re.search(r"应届(?:毕业生)?", line)), "")
        experience_source = "explicit" if experience else "unknown"
    name = _first_match(r"(?:姓名|姓\s*名)\s*[：:]\s*([\u4e00-\u9fff·]{2,20})", text)
    if not name:
        name = next((line for line in lines[:5] if re.fullmatch(r"[\u4e00-\u9fff·]{2,4}", line)
                     and not _heading(line)), "")
    phone_match = re.search(r"(?<!\d)(1[3-9]\d)[ -]?(\d{4})[ -]?(\d{4})(?!\d)", text)
    phone = "".join(phone_match.groups()) if phone_match else ""
    email = _first_match(r"([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})", text, re.IGNORECASE)
    target_job = _first_match(r"(?:目标岗位|求职意向|应聘岗位)\s*[：:]\s*([^\n；;]{2,40})", text)
    target_job_source = "explicit" if target_job else "unknown"
    if not target_job:
        target_job, target_job_source = _infer_target_job(text, filename)
    degree = _first_match(r"(博士|硕士|本科|大专|专科)", education or text)
    major = _first_match(r"([\u4e00-\u9fffA-Za-z0-9+#.]{2,24})专业", education)
    if not major:
        major = _first_match(r"(?:专业|主修)\s*[：:]\s*([^\n，,；;]{2,30})", text)
    return dict(name=name, phone=phone, email=email, target_job=target_job, target_job_source=target_job_source,
                education=education, degree=degree, major=major, experience=experience,
                experience_source=experience_source, work_experience=work,
                projects=projects, skills_raw=skills)


def extract_resume_document(filename: str, data: bytes) -> ResumeDocumentExtractResult:
    safe_name = Path(filename or "resume").name
    suffix = Path(safe_name).suffix.lower()
    if suffix == ".doc":
        raise ResumeDocumentError("暂不支持旧版 .doc，请在 Word 中另存为 .docx 后上传。")
    if suffix not in SUPPORTED:
        raise ResumeDocumentError("仅支持 PDF、DOCX 和 TXT 格式。")
    if not data:
        raise ResumeDocumentError("上传文件为空。")
    if len(data) > MAX_FILE_BYTES:
        raise ResumeDocumentError("文件不能超过 8MB。")
    raw, warnings = {".pdf": _extract_pdf, ".docx": _extract_docx, ".txt": _extract_txt}[suffix](data)
    text = _clean_text(raw)
    if not text:
        if warnings:
            return ResumeDocumentExtractResult(file_name=safe_name, file_type=SUPPORTED[suffix],
                                               raw_text="", character_count=0, warnings=warnings)
        raise ResumeDocumentError("没有从文件中提取到可用文本。")
    return ResumeDocumentExtractResult(file_name=safe_name, file_type=SUPPORTED[suffix], raw_text=text,
                                       character_count=len(text), warnings=warnings, **_draft_fields(text, safe_name))
